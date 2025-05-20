import email
from email.header import decode_header
import docx
from bs4 import BeautifulSoup
from datetime import datetime
import re
import os
import sys
import argparse
from pathlib import Path

class EMLToDOCXConverter:
    def __init__(self, eml_file_path, docx_file_path):
        self.eml_file_path = eml_file_path
        self.docx_file_path = docx_file_path
        self.doc = docx.Document()

    def decode_str(self, s):
        """安全地解码字符串"""
        if not s:
            return ""
        try:
            if isinstance(s, bytes):
                s = s.decode('utf-8', errors='ignore')
            decoded = decode_header(s)
            value, charset = decoded[0]
            if charset:
                return value.decode(charset, errors='ignore')
            return str(value)
        except Exception:
            return str(s)

    def parse_date(self, date_str):
        """解析邮件日期，返回datetime对象"""
        try:
            # 尝试多种日期格式
            date_patterns = [
                "%a, %d %b %Y %H:%M:%S %z",  # RFC 2822
                "%a, %d %b %Y %H:%M:%S %Z",  # RFC 2822 with timezone name
                "%d %b %Y %H:%M:%S %z",      # RFC 2822 without weekday
                "%d %b %Y %H:%M:%S %Z",      # RFC 2822 without weekday and with timezone name
                "%Y-%m-%d %H:%M:%S",         # ISO 8601
                "%Y-%m-%d %H:%M:%S %z",      # ISO 8601 with timezone
            ]
            
            for pattern in date_patterns:
                try:
                    return datetime.strptime(date_str, pattern)
                except ValueError:
                    continue
            
            # 如果所有格式都不匹配，尝试使用email.utils.parsedate_to_datetime
            from email.utils import parsedate_to_datetime
            return parsedate_to_datetime(date_str)
        except Exception:
            return None

    def extract_email_content(self, msg):
        """提取邮件内容"""
        subject = self.decode_str(msg.get("Subject", ""))
        from_ = self.decode_str(msg.get("From", ""))
        to = self.decode_str(msg.get("To", ""))
        date_str = self.decode_str(msg.get("Date", ""))
        date = self.parse_date(date_str)

        content = ""
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == "text/plain":
                    try:
                        body = part.get_payload(decode=True)
                        if isinstance(body, bytes):
                            content += body.decode('utf-8', errors='ignore')
                        else:
                            content += str(body)
                    except Exception:
                        continue
                elif content_type == "text/html":
                    try:
                        body = part.get_payload(decode=True)
                        if isinstance(body, bytes):
                            body = body.decode('utf-8', errors='ignore')
                        soup = BeautifulSoup(body, "html.parser")
                        content += soup.get_text()
                        tables = soup.find_all("table")
                        for table in tables:
                            rows = table.find_all("tr")
                            doc_table = self.doc.add_table(rows=len(rows), cols=len(rows[0].find_all(["th", "td"])))
                            for i, row in enumerate(rows):
                                cells = row.find_all(["th", "td"])
                                for j, cell in enumerate(cells):
                                    doc_table.cell(i, j).text = cell.get_text()
                    except Exception:
                        continue
        else:
            try:
                body = msg.get_payload(decode=True)
                if isinstance(body, bytes):
                    content = body.decode('utf-8', errors='ignore')
                else:
                    content = str(body)
            except Exception:
                content = str(msg.get_payload())

        return subject, from_, to, date, content

    def get_emails_from_eml(self, raw_content):
        """从原始EML内容中提取所有邮件"""
        # 使用正则表达式分割邮件
        email_separator = r"\n\nFrom "
        parts = re.split(email_separator, raw_content.decode('utf-8', errors='ignore'))
        
        emails = []
        for part in parts:
            if not part.strip():
                continue
            # 添加From行
            if not part.startswith("From"):
                part = "From " + part
            
            try:
                msg = email.message_from_string(part)
                subject, from_, to, date, content = self.extract_email_content(msg)
                if date:  # 只添加有有效日期的邮件
                    emails.append({
                        'subject': subject,
                        'from_': from_,
                        'to': to,
                        'date': date,
                        'content': content
                    })
            except Exception as e:
                print(f"解析邮件时出错: {e}")
                continue
        
        return emails

    def convert(self):
        try:
            with open(self.eml_file_path, "rb") as f:
                raw_content = f.read()
                
            # 获取所有邮件并按日期排序
            emails = self.get_emails_from_eml(raw_content)
            emails.sort(key=lambda x: x['date'])
            
            # 将所有邮件写入文档
            for email_data in emails:
                self.doc.add_heading(email_data['subject'], level=1)
                self.doc.add_paragraph(f"发件人: {email_data['from_']}")
                self.doc.add_paragraph(f"收件人: {email_data['to']}")
                self.doc.add_paragraph(f"日期: {email_data['date'].strftime('%Y-%m-%d %H:%M:%S')}")
                self.doc.add_paragraph(email_data['content'])
                self.doc.add_paragraph("\n")  # 添加空行分隔邮件
            
            self.doc.save(self.docx_file_path)
            print(f"转换完成，已保存到 {self.docx_file_path}")
        except Exception as e:
            print(f"转换过程中出现错误: {e}")

def convert_eml_to_docx(eml_path, docx_path):
    """Convert a single EML file to DOCX"""
    try:
        converter = EMLToDOCXConverter(eml_path, docx_path)
        converter.convert()
        print(f"Successfully converted {eml_path} to {docx_path}")
    except Exception as e:
        print(f"Error converting {eml_path}: {str(e)}")

def process_eml_files(eml_file):
    """Process all EML files in the specified directory"""
    # eml_files = [f for f in Path(directory).glob('*.eml')]
    # if not eml_files:
    #     print(f"No EML files found in {directory}")
    #     return

    # for eml_file in eml_files:
    docx_file = eml_file.with_suffix('.docx')
    convert_eml_to_docx(str(eml_file), str(docx_file))

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Convert EML files to DOCX')
    parser.add_argument('path', nargs='?', help='Path to EML file or directory containing EML files')
    args = parser.parse_args()

    if args.path:
        path = Path(args.path)
        if path.is_file() and path.suffix.lower() == '.eml':
            # Single EML file
            docx_path = path.with_suffix('.docx')
            convert_eml_to_docx(str(path), str(docx_path))
        elif path.is_dir():
            # Directory containing EML files
            process_eml_files(str(path))
        else:
            print(f"Error: {args.path} is not a valid EML file or directory")
            sys.exit(1)
    else:
        # No arguments - process current directory
        process_eml_files('.')